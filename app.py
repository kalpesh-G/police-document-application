from flask import Flask, render_template, request, redirect, url_for, send_file, flash, after_this_request, session
import json
import os
import zipfile
from docx import Document
from docxcompose.composer import Composer
from docx2pdf import convert
from datetime import datetime
import subprocess

app = Flask(__name__)
app.secret_key = 'police_gujarat_secure_key'


# # Configuration
# DATA_FILE = 'data/case_data.json'
# TEMPLATE_DIR = 'word_templates'
# GENERATED_DIR = 'generated_docs'

# # Ensure directories exist
# os.makedirs(GENERATED_DIR, exist_ok=True)
# os.makedirs('data', exist_ok=True)

# Configuration
import tempfile
TEMPLATE_DIR = 'word_templates'
GENERATED_DIR = tempfile.gettempdir()

# --------------------------------------------------
# DOCUMENT MAPPING LOGIC
# --------------------------------------------------
OFFENCE_MAPPING = {
    "281": [
        "Arrest_Memo_281.docx",
        "Bail_Bond_281.docx",
        "Notice_281.docx"
    ],
    "302": [
        "Arrest_Memo_Major.docx",
        "Remand_Application.docx",
        "Panchnama_Scene.docx"
    ],
    "379": [
        "Arrest_Memo_Theft.docx",
        "Recovery_Panchnama.docx"
    ],
    "GENERAL": [
        "Standard_Intimation.docx"
    ]
}

def get_template_path(filename, section=None):
    """Get full path to template file, using section if provided"""
    if section and section in OFFENCE_MAPPING:
        return os.path.join(TEMPLATE_DIR, section, filename)
    else:
        # Try to find which section this file belongs to
        for sec in OFFENCE_MAPPING:
            if filename in OFFENCE_MAPPING[sec]:
                return os.path.join(TEMPLATE_DIR, sec, filename)
        # Fallback to root (for backwards compatibility)
        return os.path.join(TEMPLATE_DIR, filename)

# --------------------------------------------------
# DATA VARIABLES LIST (For validation and cleaning)
# --------------------------------------------------
REQUIRED_FIELDS = [
    # Accused
    "acc_name", "acc_father", "acc_surname", "acc_alias",
    "acc_gender", "acc_age", "acc_dob",
    "acc_religion", "acc_caste", "acc_subcaste", "acc_nationality", "acc_marital",
    # Occupation
    "occ_type", "occ_place", "occ_income",
    # Perm Address
    "perm_house", "perm_area", "perm_village", "perm_district", "perm_taluka", "perm_state", "perm_pin",
    # Curr Address
    "curr_address", "curr_city", "curr_district", "curr_taluka", "curr_state", "curr_pin",
    # Contact & ID
    "mobile_1", "mobile_2", "id_type", "id_number",
    # Physical
    "phy_height", "phy_build", "phy_complexion", "phy_eyes", "phy_hair", "phy_facial_hair",
    "mark_1", "mark_2", "old_wounds", "other_id_marks",
    # Relative
    "rel_name", "rel_relation", "rel_mobile", "rel_address",
    # Case
    "case_ps", "case_district", "case_taluka",
    "crime_no", "crime_type", "crime_year",
    "offence_desc", "offence_section",
    "offence_place", "offence_date", "offence_time",
    # Arrest
    "is_arrested", "arrest_date", "arrest_time", "arrest_place", "arrest_entry_no",
    "is_bailed", "bail_authority", "bail_date", "bail_time", "bail_conditions",
    # Status
    "status_proven", "status_chargesheet", "status_case", "status_release_date",
    "intimation_method", "intimation_date", "intimation_time", "intimation_entry_no", "intimation_entry_time",
    # Officer
    "remarks", "io_name", "io_designation", "io_buckle", "io_police_station",
    "auth_date", "auth_place", "auth_print_name"
]

# --------------------------------------------------
# HELPER FUNCTIONS
# --------------------------------------------------
def load_data():
    if 'case_data' in session:
        return session['case_data']
    return {}

@app.before_request
def sync_session():
    if 'case_data' not in session:
        session['case_data'] = {}

def save_data(data):
    session['case_data'] = data
    session.modified = True

def replace_text_in_paragraph(paragraph, data):
    if not paragraph.runs:
        return

    for key in REQUIRED_FIELDS:
        placeholder = f"[{key}]"
        while True:
            full_text = "".join(run.text for run in paragraph.runs)
            idx = full_text.find(placeholder)
            if idx == -1:
                break

            start = idx
            end = idx + len(placeholder)
            value = str(data.get(key, ""))

            run_positions = []
            pos = 0
            for i, run in enumerate(paragraph.runs):
                run_positions.append((i, pos, pos + len(run.text)))
                pos += len(run.text)

            start_run_idx = None
            end_run_idx = None
            start_offset = 0
            end_offset = 0
            for i, s, e in run_positions:
                if start_run_idx is None and start >= s and start < e:
                    start_run_idx = i
                    start_offset = start - s
                if end_run_idx is None and end > s and end <= e:
                    end_run_idx = i
                    end_offset = end - s
                if start_run_idx is not None and end_run_idx is not None:
                    break

            if start_run_idx is None or end_run_idx is None:
                break

            if start_run_idx == end_run_idx:
                run = paragraph.runs[start_run_idx]
                run.text = run.text[:start_offset] + value + run.text[end_offset:]
            else:
                start_run = paragraph.runs[start_run_idx]
                end_run = paragraph.runs[end_run_idx]
                start_text = start_run.text
                end_text = end_run.text
                start_run.text = start_text[:start_offset] + value + end_text[end_offset:]
                for i in range(start_run_idx + 1, end_run_idx + 1):
                    paragraph.runs[i].text = ""

def generate_document(template_name, data, output_path, section=None):
    template_path = get_template_path(template_name, section)
    if not os.path.exists(template_path):
        return False
    
    doc = Document(template_path)
    
    # Replace in paragraphs
    for p in doc.paragraphs:
        replace_text_in_paragraph(p, data)
        
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_text_in_paragraph(p, data)
                    
    doc.save(output_path)
    return True

def generate_pdf_from_docx(docx_path, pdf_path):
    try:
        convert(docx_path, pdf_path)
        if os.path.exists(pdf_path):
            return True
    except Exception:
        pass

    try:
        out_dir = os.path.dirname(pdf_path)
        subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf", "--outdir", out_dir, docx_path],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE
        )
        return os.path.exists(pdf_path)
    except Exception:
        return False

def extract_placeholders_from_docx(template_path):
    if not os.path.exists(template_path):
        return set()

    doc = Document(template_path)
    found = set()

    def scan_text(text):
        for key in REQUIRED_FIELDS:
            placeholder = f"[{key}]"
            if placeholder in text:
                found.add(key)

    for p in doc.paragraphs:
        scan_text(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    scan_text(p.text)

    return found

def extract_preview_from_docx(template_path, data):
    if not os.path.exists(template_path):
        return ""

    doc = Document(template_path)
    html_content = "<style>"
    html_content += """
    .docx-preview-para { margin: 0.5em 0; line-height: 1.4; }
    .docx-preview-table { width: 100%; margin: 0.5em 0; border-collapse: collapse; border: 1px solid #999; }
    .docx-preview-table td { border: 1px solid #999; padding: 8px; vertical-align: top; }
    .docx-preview-table th { border: 1px solid #999; padding: 8px; background-color: #e9ecef; font-weight: bold; }
    .docx-preview-run { display: inline; }
    </style>"""

    for p in doc.paragraphs:
        # Create a copy and replace placeholders
        p_copy = Document().add_paragraph()
        for run in p.runs:
            p_copy.add_run(run.text)
        replace_text_in_paragraph(p_copy, data)

        # Build HTML with formatting
        p_html = "<p class='docx-preview-para' style='"
        
        # Paragraph alignment
        alignment = p.alignment
        if alignment == 1:  # CENTER
            p_html += "text-align: center; "
        elif alignment == 2:  # RIGHT
            p_html += "text-align: right; "
        elif alignment == 3:  # JUSTIFY
            p_html += "text-align: justify; "
        
        # Paragraph indentation/spacing
        if p.paragraph_format.left_indent:
            p_html += f"margin-left: {p.paragraph_format.left_indent.pt}pt; "
        if p.paragraph_format.right_indent:
            p_html += f"margin-right: {p.paragraph_format.right_indent.pt}pt; "
        if p.paragraph_format.space_before:
            p_html += f"margin-top: {p.paragraph_format.space_before.pt}pt; "
        if p.paragraph_format.space_after:
            p_html += f"margin-bottom: {p.paragraph_format.space_after.pt}pt; "
        
        p_html += "'>"

        # Add runs with formatting
        for run in p_copy.runs:
            run_html = "<span class='docx-preview-run' style='"
            
            # Font properties
            if run.font.size:
                run_html += f"font-size: {run.font.size.pt}pt; "
            if run.font.bold:
                run_html += "font-weight: bold; "
            if run.font.italic:
                run_html += "font-style: italic; "
            if run.font.underline:
                run_html += "text-decoration: underline; "
            if run.font.color and run.font.color.rgb:
                run_html += f"color: #{str(run.font.color.rgb)}; "
            
            run_html += "'>" + (run.text or "") + "</span>"
            p_html += run_html

        p_html += "</p>"
        
        # Only add non-empty paragraphs
        if p_copy.text.strip():
            html_content += p_html
        else:
            # Add empty paragraph for spacing
            html_content += "<p class='docx-preview-para' style='height: 0.5em;'></p>"

    # Process tables
    for table in doc.tables:
        html_content += "<table class='docx-preview-table'>"
        
        for row_idx, row in enumerate(table.rows):
            html_content += "<tr>"
            
            for cell in row.cells:
                # Check if this is a header cell (first row)
                is_header = row_idx == 0
                tag = "th" if is_header else "td"
                
                cell_html = f"<{tag}>"
                
                for p in cell.paragraphs:
                    # Create copy and replace placeholders
                    p_copy = Document().add_paragraph()
                    for run in p.runs:
                        p_copy.add_run(run.text)
                    replace_text_in_paragraph(p_copy, data)
                    
                    # Build cell paragraph
                    cell_p_html = "<div style='"
                    if p.alignment == 1:  # CENTER
                        cell_p_html += "text-align: center; "
                    elif p.alignment == 2:  # RIGHT
                        cell_p_html += "text-align: right; "
                    cell_p_html += "'>"
                    
                    # Add runs
                    for run in p_copy.runs:
                        run_text = run.text or ""
                        if run.font.bold:
                            run_text = f"<b>{run_text}</b>"
                        if run.font.italic:
                            run_text = f"<i>{run_text}</i>"
                        cell_p_html += run_text
                    
                    cell_p_html += "</div>"
                    cell_html += cell_p_html
                
                cell_html += f"</{tag}>"
                html_content += cell_html
            
            html_content += "</tr>"
        
        html_content += "</table>"

    return html_content

# --------------------------------------------------
# API ENDPOINTS FOR LOCALSTORAGE SYNC
# --------------------------------------------------

@app.route('/api/save_data', methods=['POST'])
def api_save_data():
    data = request.get_json() or {}
    save_data(data)
    return {'status': 'success'}, 200

@app.route('/api/load_data', methods=['GET'])
def api_load_data():
    data = load_data()
    return data, 200

# --------------------------------------------------
# ROUTES
# --------------------------------------------------

@app.route('/', methods=['GET', 'POST'])
def home():
    data = load_data()
    if request.method == 'POST':
        # Update data with form submission
        for field in REQUIRED_FIELDS:
            data[field] = request.form.get(field, "")
        
        save_data(data)
        
        # If user clicked "Proceed", go to documents page
        if 'proceed' in request.form:
            return redirect(url_for('documents'))
            
    return render_template('index.html', data=data, fields=REQUIRED_FIELDS)

@app.route('/documents', methods=['GET', 'POST'])
def documents():
    data = load_data()
    
    # 1. Update data if missing fields form is submitted
    if request.method == 'POST':
        for key in request.form:
            if key in REQUIRED_FIELDS:
                data[key] = request.form[key]
        save_data(data)

    # 2. Determine Documents based on Offence Section
    section = data.get('offence_section', 'GENERAL')
    # Default to 'GENERAL' if section not in map
    doc_list = OFFENCE_MAPPING.get(section, OFFENCE_MAPPING['GENERAL'])
    
    # 3. Check for Missing Data (only placeholders used in selected templates)
    required_in_docs = set()
    for filename in doc_list:
        template_path = get_template_path(filename, section)
        required_in_docs.update(extract_placeholders_from_docx(template_path))

    missing_fields = [f for f in REQUIRED_FIELDS if f in required_in_docs and not data.get(f)]
    
    return render_template('documents.html', 
                           data=data, 
                           docs=doc_list, 
                           missing_fields=missing_fields,
                           section=section)

@app.route('/download_single/<filename>')
def download_single(filename):
    data = load_data()
    section = data.get('offence_section', 'GENERAL')
    output_path = os.path.join(GENERATED_DIR, f"Filled_{filename}")
    
    success = generate_document(filename, data, output_path, section)
    if success:
        @after_this_request
        def cleanup(response):
            try:
                os.remove(output_path)
            except:
                pass
            return response
        return send_file(output_path, as_attachment=True)
    return "Template not found", 404

@app.route('/download_single_pdf/<filename>')
def download_single_pdf(filename):
    data = load_data()
    section = data.get('offence_section', 'GENERAL')
    output_docx = os.path.join(GENERATED_DIR, f"Filled_{filename}")
    output_pdf = os.path.join(GENERATED_DIR, f"Filled_{os.path.splitext(filename)[0]}.pdf")

    success = generate_document(filename, data, output_docx, section)
    if not success:
        return "Template not found", 404

    if not generate_pdf_from_docx(output_docx, output_pdf):
        return "PDF generation failed", 500

    @after_this_request
    def cleanup(response):
        try:
            os.remove(output_docx)
            os.remove(output_pdf)
        except:
            pass
        return response
    return send_file(output_pdf, as_attachment=True)

@app.route('/download_all_zip')
def download_all_zip():
    data = load_data()
    section = data.get('offence_section', 'GENERAL')
    doc_list = OFFENCE_MAPPING.get(section, OFFENCE_MAPPING['GENERAL'])
    
    zip_path = os.path.join(GENERATED_DIR, "All_Documents.zip")
    generated_files = []
    with zipfile.ZipFile(zip_path, 'w') as zipf:
        for filename in doc_list:
            output_name = f"Filled_{filename}"
            output_path = os.path.join(GENERATED_DIR, output_name)
            if generate_document(filename, data, output_path, section):
                zipf.write(output_path, arcname=output_name)
                generated_files.append(output_path)
    
    @after_this_request
    def cleanup(response):
        try:
            os.remove(zip_path)
            for f in generated_files:
                try:
                    os.remove(f)
                except:
                    pass
        except:
            pass
        return response
    return send_file(zip_path, as_attachment=True)

@app.route('/download_all_pdf_zip')
def download_all_pdf_zip():
    data = load_data()
    section = data.get('offence_section', 'GENERAL')
    doc_list = OFFENCE_MAPPING.get(section, OFFENCE_MAPPING['GENERAL'])

    zip_path = os.path.join(GENERATED_DIR, "All_Documents_PDF.zip")
    generated_files = []
    with zipfile.ZipFile(zip_path, 'w') as zipf:
        for filename in doc_list:
            output_docx = os.path.join(GENERATED_DIR, f"Filled_{filename}")
            output_pdf = os.path.join(GENERATED_DIR, f"Filled_{os.path.splitext(filename)[0]}.pdf")
            if generate_document(filename, data, output_docx, section):
                if generate_pdf_from_docx(output_docx, output_pdf):
                    zipf.write(output_pdf, arcname=os.path.basename(output_pdf))
                    generated_files.extend([output_docx, output_pdf])

    @after_this_request
    def cleanup(response):
        try:
            os.remove(zip_path)
            for f in generated_files:
                try:
                    os.remove(f)
                except:
                    pass
        except:
            pass
        return response
    return send_file(zip_path, as_attachment=True)

@app.route('/download_merged')
def download_merged():
    data = load_data()
    section = data.get('offence_section', 'GENERAL')
    doc_list = OFFENCE_MAPPING.get(section, OFFENCE_MAPPING['GENERAL'])
    
    if not doc_list:
        return "No documents to merge", 400

    # Generate first document as base
    base_path = os.path.join(GENERATED_DIR, "Merged_Master.docx")
    generate_document(doc_list[0], data, base_path, section)
    
    master_doc = Document(base_path)
    composer = Composer(master_doc)
    
    # Append subsequent documents
    temp_files = []
    for filename in doc_list[1:]:
        temp_path = os.path.join(GENERATED_DIR, f"temp_{filename}")
        generate_document(filename, data, temp_path, section)
        doc_to_append = Document(temp_path)
        master_doc.add_page_break()
        composer.append(doc_to_append)
        temp_files.append(temp_path)
        
    composer.save(base_path)
    
    @after_this_request
    def cleanup(response):
        try:
            os.remove(base_path)
            for f in temp_files:
                try:
                    os.remove(f)
                except:
                    pass
        except:
            pass
        return response
    return send_file(base_path, as_attachment=True)

@app.route('/download_merged_pdf')
def download_merged_pdf():
    data = load_data()
    section = data.get('offence_section', 'GENERAL')
    doc_list = OFFENCE_MAPPING.get(section, OFFENCE_MAPPING['GENERAL'])

    if not doc_list:
        return "No documents to merge", 400

    base_path = os.path.join(GENERATED_DIR, "Merged_Master.docx")
    generate_document(doc_list[0], data, base_path, section)

    master_doc = Document(base_path)
    composer = Composer(master_doc)

    temp_files = []
    for filename in doc_list[1:]:
        temp_path = os.path.join(GENERATED_DIR, f"temp_{filename}")
        generate_document(filename, data, temp_path, section)
        doc_to_append = Document(temp_path)
        master_doc.add_page_break()
        composer.append(doc_to_append)
        temp_files.append(temp_path)

    composer.save(base_path)

    merged_pdf = os.path.join(GENERATED_DIR, "Merged_Master.pdf")
    if not generate_pdf_from_docx(base_path, merged_pdf):
        return "PDF generation failed", 500

    @after_this_request
    def cleanup(response):
        try:
            os.remove(base_path)
            os.remove(merged_pdf)
            for f in temp_files:
                try:
                    os.remove(f)
                except:
                    pass
        except:
            pass
        return response
    return send_file(merged_pdf, as_attachment=True)

@app.route('/preview/<filename>')
def preview_document(filename):
    data = load_data()
    section = data.get('offence_section', 'GENERAL')
    template_path = get_template_path(filename, section)
    
    if not os.path.exists(template_path):
        return "Template not found", 404
    
    preview_html = extract_preview_from_docx(template_path, data)
    return preview_html

@app.route('/preview_merged')
def preview_merged():
    data = load_data()
    section = data.get('offence_section', 'GENERAL')
    doc_list = OFFENCE_MAPPING.get(section, OFFENCE_MAPPING['GENERAL'])
    
    merged_html = ""
    for filename in doc_list:
        template_path = get_template_path(filename, section)
        if os.path.exists(template_path):
            merged_html += f"<div style='page-break-after: always; margin-bottom: 30px; padding-bottom: 20px; border-bottom: 3px solid #ccc;'>"
            merged_html += f"<h4 style='color: #003366; margin-bottom: 15px;'>📄 {filename}</h4>"
            merged_html += extract_preview_from_docx(template_path, data)
            merged_html += "</div>"
    
    if not merged_html:
        return "<p class='text-muted'>No documents to preview</p>"
    
    return merged_html

if __name__ == '__main__':
    app.run(debug=True, port=5000)