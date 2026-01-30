# setup_project.py
import os
from docx import Document

folders = ['templates', 'static', 'word_templates', 'generated_docs', 'data']
for folder in folders:
    os.makedirs(folder, exist_ok=True)

# Define sample templates with placeholders
templates_data = {
    "arrest_memo.docx": "ARREST MEMO\n\nAccused: [acc_name]\nAge: [acc_age]\nOffence: [offence_section]\nDate: [offence_date]\nOfficer: [io_name]",
    "notice_41a.docx": "NOTICE UNDER 41A\n\nTo, [acc_name]\nAddress: [curr_address]\n\nPlease appear before [io_name] at [io_police_station].\nCase: [crime_no]",
    "bail_bond.docx": "BAIL BOND\n\nI, [acc_name], son of [acc_father], caste [acc_caste], hereby agree to terms.\nWitness: [rel_name]\nAmount: [occ_income]",
    "panchnama.docx": "PANCHNAMA\n\nPlace: [offence_place]\nDescription: [offence_desc]\nRecovered from: [acc_name]\nMarkings: [mark_1], [mark_2]"
}

print("Generating sample templates...")
for filename, content in templates_data.items():
    path = os.path.join('word_templates', filename)
    if not os.path.exists(path):
        doc = Document()
        doc.add_heading(filename.replace('.docx', '').upper(), 0)
        doc.add_paragraph(content)
        doc.save(path)
        print(f"Created: {path}")

print("\nSetup Complete! Now run 'python app.py'")